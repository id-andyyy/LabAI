#!/usr/bin/env python3
"""Генератор ГОСТ-совместимого Word-документа из report.md.

Два режима работы:
  1. Без шаблона — создаёт документ с нуля (титульный лист, стили, содержимое).
  2. С шаблоном-пустышкой (--template) — открывает готовый шаблон с титульным листом
     и заголовками секций (Цели работы, Ход работы, и т.д.), вставляет содержимое
     между секциями, копирует стиль заголовков из шаблона.

Аргументы:
    --config PATH      config.json (обязательно)
    --report PATH      report.md (обязательно)
    --output PATH      Выходной файл .docx (обязательно)
    --images PATH      Директория с изображениями (опционально)
    --image-map PATH   image_map.json — маппинг номер→файл (опционально)
    --template PATH    Шаблон-пустышка .docx с титульным листом и заголовками
    --no-title-page    Не создавать титульный лист (режим без шаблона)
"""

import argparse
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


# ---------------------------------------------------------------------------
# Константы ГОСТ
# ---------------------------------------------------------------------------
FONT_NAME = "Times New Roman"
FONT_SIZE_NORMAL = Pt(14)
FONT_SIZE_HEADING1 = Pt(16)
FONT_SIZE_HEADING2 = Pt(14)
FONT_SIZE_CAPTION = Pt(14)  # Такой же размер, как основной текст
FONT_SIZE_TITLE = Pt(16)
LINE_SPACING = 1.5
LINE_SPACING_SINGLE = 1.0
FIRST_LINE_INDENT = Cm(1.5)
PAGE_MARGIN = Cm(2)
# Макс. ширина изображения = ширина текста минус красная строка: 17 - 1.5 = 15.5 см
MAX_IMAGE_WIDTH = Cm(15.5)

# ---------------------------------------------------------------------------
# Регулярные выражения
# ---------------------------------------------------------------------------

# Нумерованные заголовки
RE_HEADING1 = re.compile(r"^(\d+)\.\s+(.+)$")
RE_HEADING2 = re.compile(r"^(\d+\.\d+)\.?\s+(.+)$")
RE_HEADING3 = re.compile(r"^(\d+\.\d+\.\d+)\.?\s+(.+)$")

# Markdown-заголовки
RE_MD_H0 = re.compile(r"^#\s+(.+)$")         # # Заголовок  -> H1
RE_MD_H1 = re.compile(r"^##\s+(.+)$")        # ## Заголовок -> H1
RE_MD_H2 = re.compile(r"^###\s+(.+)$")       # ### Подзаголовок -> H2
RE_MD_H3 = re.compile(r"^####\s+(.+)$")      # #### Подподзаголовок -> H3

# Плейсхолдеры и подписи
RE_PLACEHOLDER = re.compile(r"\[ВСТАВИТЬ РИСУНОК (\d+) ЗДЕСЬ\]")
RE_FIGURE_CAPTION = re.compile(r"^Рисунок (\d+)\.\s+(.+)$")
RE_TABLE_CAPTION = re.compile(r"^Таблица (\d+)\.\s+(.+)$")

# Инлайн-форматирование
RE_BOLD = re.compile(r"\*\*(.+?)\*\*")
RE_ITALIC = re.compile(r"\*(.+?)\*")
RE_CODE_FENCE = re.compile(r"^```")
RE_INLINE_CODE = re.compile(r"`([^`]+)`")

# Markdown-таблицы
RE_TABLE_ROW = re.compile(r"^\|(.+)\|$")
RE_TABLE_SEP = re.compile(r"^\|[\s:]*-+[\s:|-]*\|$")

# Списки (нумерованные используют только скобку: 1) 2) 3) — по ГОСТ)
RE_LIST_BULLET = re.compile(r"^[-*]\s+(.+)$")
RE_LIST_NUMBERED = re.compile(r"^(\d+)\)\s+(.+)$")

# Известные названия секций шаблона (ключ -> варианты текста)
KNOWN_SECTIONS = {
    "цели работы": ["цели работы", "цель работы"],
    "ход работы": ["ход работы", "ход выполнения работы"],
    "дополнительное исследование": ["дополнительное исследование"],
    "выводы": ["выводы", "вывод"],
}

# Порядок обработки секций
SECTION_ORDER = ["цели работы", "ход работы", "дополнительное исследование", "выводы"]


# ---------------------------------------------------------------------------
# Вспомогательные функции для работы с шаблоном
# ---------------------------------------------------------------------------
def _insert_para_after(ref_para, doc):
    """Вставить пустой параграф после ref_para, вернуть Paragraph."""
    new_p = OxmlElement('w:p')
    ref_para._element.addnext(new_p)
    # _body (BlockItemContainer) предоставляет .part для работы со стилями
    parent = doc._body if hasattr(doc, '_body') else doc.element.body
    return Paragraph(new_p, parent)


def _normalize_section_name(text):
    """Нормализовать название секции для сопоставления."""
    text = text.strip().lower()
    text = re.sub(r'^\d+\.\s*', '', text)
    return text


def _match_section_key(name):
    """Сопоставить название секции из report.md с ключом шаблона."""
    normalized = _normalize_section_name(name)
    for key, variants in KNOWN_SECTIONS.items():
        if normalized in variants or normalized == key:
            return key
    return None


def find_template_sections(doc):
    """Найти параграфы-заголовки секций в шаблоне.

    Returns:
        dict: {section_key: paragraph_object}
    """
    found = {}
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        normalized = _normalize_section_name(text)
        for key, variants in KNOWN_SECTIONS.items():
            if normalized in variants and key not in found:
                found[key] = para
                break
    return found


def _clone_paragraph_format(source_para, target_para):
    """Скопировать форматирование параграфа (отступы, интервалы, выравнивание)."""
    sf = source_para.paragraph_format
    tf = target_para.paragraph_format
    if sf.alignment is not None:
        tf.alignment = sf.alignment
    if sf.space_before is not None:
        tf.space_before = sf.space_before
    if sf.space_after is not None:
        tf.space_after = sf.space_after
    if sf.line_spacing is not None:
        tf.line_spacing = sf.line_spacing
    if sf.first_line_indent is not None:
        tf.first_line_indent = sf.first_line_indent


def _clone_run_format(source_run, target_run):
    """Скопировать форматирование символов из source_run в target_run."""
    if source_run.font.name:
        target_run.font.name = source_run.font.name
    if source_run.font.size:
        target_run.font.size = source_run.font.size
    if source_run.font.bold is not None:
        target_run.font.bold = source_run.font.bold
    if source_run.font.italic is not None:
        target_run.font.italic = source_run.font.italic
    if source_run.font.color and source_run.font.color.rgb:
        target_run.font.color.rgb = source_run.font.color.rgb


def _set_contextual_spacing(para_or_element):
    """Установить флаг 'Не добавлять интервал между абзацами одного стиля'.

    Принимает Paragraph-объект или lxml-элемент стиля.
    """
    el = para_or_element._element if hasattr(para_or_element, '_element') else para_or_element
    pPr = el.get_or_add_pPr()
    cs = OxmlElement('w:contextualSpacing')
    pPr.append(cs)


def _is_heading1(line):
    """Проверить, является ли строка заголовком H1 (секции).

    Returns:
        str или None — текст заголовка без номера/хеша.
    """
    m = RE_HEADING1.match(line)
    if m:
        return m.group(2).strip()
    m = RE_MD_H1.match(line)
    if m:
        return m.group(1).strip()
    m = RE_MD_H0.match(line)
    if m:
        return m.group(1).strip()
    return None


def parse_report_into_sections(report_text):
    """Разбить report.md на секции по главным заголовкам (H1).

    Поддерживаются форматы: `1. Заголовок`, `## Заголовок`, `# Заголовок`.

    Returns:
        list of tuples: [(section_key, lines), ...]
    """
    sections = []
    current_key = None
    current_lines = []

    for line in report_text.split('\n'):
        stripped = line.rstrip()
        heading_text = _is_heading1(stripped)

        if heading_text is not None:
            if current_key is not None:
                sections.append((current_key, current_lines))
            matched = _match_section_key(heading_text)
            current_key = matched if matched else heading_text.lower()
            current_lines = []
        else:
            current_lines.append(stripped)

    if current_key is not None:
        sections.append((current_key, current_lines))

    return sections


# ---------------------------------------------------------------------------
# Таблицы
# ---------------------------------------------------------------------------

def _collect_table_lines(lines, start):
    """Собрать строки markdown-таблицы начиная с start.

    Returns:
        (table_lines, end_index)
    """
    table_lines = []
    i = start
    while i < len(lines):
        line = lines[i].rstrip() if isinstance(lines[i], str) else lines[i]
        if RE_TABLE_ROW.match(line) or RE_TABLE_SEP.match(line):
            table_lines.append(line)
            i += 1
        else:
            break
    return table_lines, i


def _parse_markdown_table(table_lines):
    """Парсинг markdown-таблицы.

    Returns:
        (headers, rows) или (None, None) если некорректна.
    """
    if len(table_lines) < 2:
        return None, None

    header_line = table_lines[0]
    headers = [c.strip() for c in header_line.strip('|').split('|')]

    data_start = 1
    if len(table_lines) > 1 and RE_TABLE_SEP.match(table_lines[1]):
        data_start = 2

    rows = []
    for line in table_lines[data_start:]:
        if RE_TABLE_SEP.match(line):
            continue
        row_cells = [c.strip() for c in line.strip('|').split('|')]
        while len(row_cells) < len(headers):
            row_cells.append("")
        rows.append(row_cells[:len(headers)])

    return headers, rows


def _format_table_cell(cell, text):
    """Отформатировать ячейку таблицы."""
    for para in cell.paragraphs:
        para.clear()
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.left_indent = Cm(0)
    para.paragraph_format.line_spacing = LINE_SPACING_SINGLE

    clean_text = RE_INLINE_CODE.sub(r'\1', text)

    if RE_BOLD.search(clean_text):
        parts = re.split(r"(\*\*.*?\*\*)", clean_text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = para.add_run(part[2:-2])
                run.font.bold = True
            else:
                run = para.add_run(part)
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE_NORMAL
    else:
        run = para.add_run(clean_text)
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_NORMAL

    # Padding ячейки
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side in ['top', 'bottom', 'left', 'right']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), '57')  # ~1mm
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def _create_word_table(doc, headers, rows, ref_para=None):
    """Создать Word-таблицу.

    Args:
        doc: Document
        headers: список заголовков
        rows: список строк
        ref_para: вставить после этого параграфа (режим шаблона)

    Returns:
        Table
    """
    num_cols = len(headers)
    num_rows = len(rows) + 1

    table = doc.add_table(rows=num_rows, cols=num_cols, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Ширина таблицы 100%
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

    # Заголовки (жирные)
    for j, header in enumerate(headers):
        cell = table.cell(0, j)
        _format_table_cell(cell, header)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.bold = True

    # Данные
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i + 1, j)
            _format_table_cell(cell, cell_text)

    # Режим шаблона: переместить таблицу после ref_para
    if ref_para is not None:
        ref_para._element.addnext(tbl)

    return table


# ---------------------------------------------------------------------------
# Списки
# ---------------------------------------------------------------------------
def _add_list_paragraph(doc, text, bullet_char="\u2013", number=None, ref_para=None):
    """Создать элемент списка.

    Args:
        doc: Document
        text: текст элемента
        bullet_char: символ маркера (по умолчанию – длинное тире)
        number: номер (для нумерованного списка)
        ref_para: вставить после этого параграфа (режим шаблона)

    Returns:
        Paragraph
    """
    if ref_para is not None:
        p = _insert_para_after(ref_para, doc)
    else:
        p = doc.add_paragraph()

    p.paragraph_format.line_spacing = LINE_SPACING
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(2.0)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    _set_contextual_spacing(p)

    prefix = f"{number}) " if number is not None else f"{bullet_char} "
    run = p.add_run(prefix)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_NORMAL

    add_formatted_run(p, text)
    return p


# ---------------------------------------------------------------------------
# Настройка стилей (для режима без шаблона)
# ---------------------------------------------------------------------------
def setup_styles(doc: Document):
    """Настроить стили документа по ГОСТ."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE_NORMAL
    font.color.rgb = RGBColor(0, 0, 0)

    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.first_line_indent = FIRST_LINE_INDENT
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _set_contextual_spacing(doc.styles["Normal"].element)

    # Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = FONT_NAME
    h1.font.size = FONT_SIZE_HEADING1
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.first_line_indent = Cm(0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = FONT_NAME
    h2.font.size = FONT_SIZE_HEADING2
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.first_line_indent = Cm(0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Heading 3
    h3 = doc.styles["Heading 3"]
    h3.font.name = FONT_NAME
    h3.font.size = FONT_SIZE_NORMAL
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.space_before = Pt(6)
    h3.paragraph_format.space_after = Pt(3)
    h3.paragraph_format.first_line_indent = Cm(0)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # FigureCaption
    if "FigureCaption" not in [s.name for s in doc.styles]:
        fc = doc.styles.add_style("FigureCaption", 1)  # 1 = WD_STYLE_TYPE.PARAGRAPH
    else:
        fc = doc.styles["FigureCaption"]
    fc.font.name = FONT_NAME
    fc.font.size = FONT_SIZE_CAPTION
    fc.font.color.rgb = RGBColor(0, 0, 0)
    fc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fc.paragraph_format.space_before = Pt(6)
    fc.paragraph_format.space_after = Pt(0)
    fc.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    fc.paragraph_format.line_spacing = LINE_SPACING

    # TableCaption
    if "TableCaption" not in [s.name for s in doc.styles]:
        tc = doc.styles.add_style("TableCaption", 1)
    else:
        tc = doc.styles["TableCaption"]
    tc.font.name = FONT_NAME
    tc.font.size = FONT_SIZE_CAPTION
    tc.font.color.rgb = RGBColor(0, 0, 0)
    tc.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    tc.paragraph_format.space_before = Pt(0)
    tc.paragraph_format.space_after = Pt(6)
    tc.paragraph_format.first_line_indent = Cm(0)


def setup_page(doc: Document):
    """Настроить поля страницы."""
    for section in doc.sections:
        section.top_margin = PAGE_MARGIN
        section.bottom_margin = PAGE_MARGIN
        section.left_margin = PAGE_MARGIN
        section.right_margin = PAGE_MARGIN


# ---------------------------------------------------------------------------
# Титульный лист (для режима без шаблона)
# ---------------------------------------------------------------------------
def add_title_page(doc: Document, config: dict):
    """Создать титульный лист из config.json."""
    student = config.get("student", {})
    institution = config.get("institution", {})
    course = config.get("course", {})

    # Верх: название учебного заведения
    inst_name = institution.get("name", "Учебное заведение")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = LINE_SPACING_SINGLE
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(inst_name)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_NORMAL
    run.font.bold = True

    # Пустое пространство
    for _ in range(6):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(0)
        sp.paragraph_format.first_line_indent = Cm(0)

    # Центр: название работы
    course_name = course.get("name", "Предмет")
    lab_number = course.get("lab_number", "")
    lab_title = course.get("lab_title", "Лабораторная работа")
    lab_type = course.get("lab_type", "Лабораторная работа")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(course_name.upper())
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_TITLE
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    lab_text = f"{lab_type} \u2116 {lab_number}" if lab_number else lab_type
    run = p.add_run(lab_text)
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_TITLE
    run.font.bold = True

    if lab_title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(f"\u00ab{lab_title}\u00bb")
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_TITLE
        run.font.bold = True

    # Пустое пространство
    for _ in range(6):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(0)
        sp.paragraph_format.first_line_indent = Cm(0)

    # Правый нижний угол: данные студента
    student_name = student.get("name", "Студент")
    student_group = student.get("group", "Группа")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = LINE_SPACING_SINGLE
    run = p.add_run(f"Выполнил: {student_name}")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_NORMAL

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = LINE_SPACING_SINGLE
    run = p.add_run(f"Группа: {student_group}")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_NORMAL

    # Преподаватель (если указан в конфиге)
    reviewer = course.get("reviewer")
    if reviewer:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = LINE_SPACING_SINGLE
        run = p.add_run(f"Преподаватель: {reviewer}")
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_NORMAL

    # Пустое пространство до низа
    for _ in range(4):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(0)
        sp.paragraph_format.first_line_indent = Cm(0)

    # Низ по центру: город, год
    city = institution.get("city", "Город")
    import datetime
    year = datetime.datetime.now().year

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"{city}, {year}")
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_NORMAL
    run.font.bold = True

    # Разрыв страницы
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Вставка изображения
# ---------------------------------------------------------------------------
def _calc_image_width(image_path):
    """Вычислить ширину изображения с учётом полей и красной строки."""
    from PIL import Image

    img = Image.open(image_path)
    width_px = img.size[0]
    dpi = img.info.get("dpi", (96, 96))
    dpi_x = dpi[0] if isinstance(dpi, tuple) else dpi
    if dpi_x == 0:
        dpi_x = 96
    width_cm = width_px / dpi_x * 2.54
    # Макс. ширина = ширина текста минус красная строка = 17 - 1.5 = 15.5 см
    if width_cm > 15.5:
        width_cm = 15.5
    return width_cm


def insert_image(doc: Document, image_path: str):
    """Вставить изображение с масштабированием (добавляет в конец документа)."""
    width_cm = _calc_image_width(image_path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = LINE_SPACING
    _set_contextual_spacing(p)
    run = p.add_run()
    run.add_picture(image_path, width=Cm(width_cm))


def _insert_image_at(doc, ref_para, image_path):
    """Вставить изображение после ref_para (для режима шаблона). Вернуть параграф."""
    width_cm = _calc_image_width(image_path)
    new_p = _insert_para_after(ref_para, doc)
    new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    new_p.paragraph_format.left_indent = Cm(0)
    new_p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    new_p.paragraph_format.space_before = Pt(6)
    new_p.paragraph_format.space_after = Pt(0)
    new_p.paragraph_format.line_spacing = LINE_SPACING
    _set_contextual_spacing(new_p)

    run = new_p.add_run()
    run.add_picture(image_path, width=Cm(width_cm))
    return new_p


# ---------------------------------------------------------------------------
# Парсинг и добавление текста
# ---------------------------------------------------------------------------
def add_formatted_run(paragraph, text: str):
    """Добавить текст с базовым форматированием (только курсив, без жирного).

    Жирный шрифт НЕ применяется в тексте — маркеры ** снимаются,
    но bold не устанавливается. Жирный допустим только в заголовках.
    Инлайн-код (`...`) — маркеры снимаются, текст вставляется как обычный.
    """
    # Убрать инлайн-код (оставить только текст внутри)
    text = RE_INLINE_CODE.sub(r'\1', text)
    # Разбить текст на сегменты с форматированием
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            # Жирный НЕ применяется — только снимаем маркеры
            run = paragraph.add_run(part[2:-2])
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE_NORMAL
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE_NORMAL
            run.italic = True
        else:
            run = paragraph.add_run(part)
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE_NORMAL


# ---------------------------------------------------------------------------
# Определение типа элемента (вспомогательная функция)
# ---------------------------------------------------------------------------
def _match_h3(line):
    """Проверить H3: 1.1.1. Title или #### Title. Вернуть текст или None."""
    m = RE_HEADING3.match(line)
    if m:
        return f"{m.group(1)}. {m.group(2)}"
    m = RE_MD_H3.match(line)
    if m:
        return m.group(1)
    return None


def _match_h2(line):
    """Проверить H2: 1.1. Title или ### Title. Вернуть текст или None."""
    m = RE_HEADING2.match(line)
    if m:
        return f"{m.group(1)}. {m.group(2)}"
    m = RE_MD_H2.match(line)
    if m:
        return m.group(1)
    return None


def _match_h1(line):
    """Проверить H1: 1. Title или ## Title или # Title. Вернуть текст или None."""
    m = RE_HEADING1.match(line)
    if m:
        return f"{m.group(1)}. {m.group(2)}"
    m = RE_MD_H1.match(line)
    if m:
        return m.group(1)
    m = RE_MD_H0.match(line)
    if m:
        return m.group(1)
    return None


# ---------------------------------------------------------------------------
# Обработка плейсхолдера изображения
# ---------------------------------------------------------------------------
def _handle_placeholder(line, image_map, images_dir):
    """Проверить наличие плейсхолдера и вернуть данные.

    Returns:
        (fig_num, img_path) или (fig_num, None) или None
    """
    m = RE_PLACEHOLDER.search(line)
    if not m:
        return None
    fig_num = m.group(1)
    img_file = image_map.get(fig_num) or image_map.get(int(fig_num))
    if img_file and images_dir:
        img_path = images_dir / img_file if not Path(img_file).is_absolute() else Path(img_file)
        if img_path.exists():
            return fig_num, str(img_path)
        return fig_num, None
    return fig_num, None


# ---------------------------------------------------------------------------
# Обработка отчёта (режим без шаблона)
# ---------------------------------------------------------------------------
def process_report(doc: Document, report_text: str, image_map: dict, images_dir: Path | None):
    """Обработать report.md и добавить содержимое в документ (режим без шаблона)."""
    lines = report_text.split("\n")
    i = 0
    in_code_block = False

    while i < len(lines):
        line = lines[i].rstrip()

        # Блоки кода (```) — пропускаем целиком
        if RE_CODE_FENCE.match(line.strip()):
            in_code_block = not in_code_block
            i += 1
            continue
        if in_code_block:
            i += 1
            continue

        # Пустая строка — пропускаем
        if not line:
            i += 1
            continue

        # --- Markdown-таблица ---
        if RE_TABLE_ROW.match(line):
            table_lines, end_idx = _collect_table_lines(lines, i)
            headers, rows = _parse_markdown_table(table_lines)
            if headers and rows is not None:
                _create_word_table(doc, headers, rows)
            i = end_idx
            continue

        # --- Списки ---
        m = RE_LIST_BULLET.match(line)
        if m:
            _add_list_paragraph(doc, m.group(1))
            i += 1
            continue

        m = RE_LIST_NUMBERED.match(line)
        if m:
            _add_list_paragraph(doc, m.group(2), number=m.group(1))
            i += 1
            continue

        # --- Заголовки (проверяем от H3 к H1, чтобы 1.1.1. не захватился как 1.1.) ---

        h3_text = _match_h3(line)
        if h3_text:
            p = doc.add_paragraph(style="Heading 3")
            p.add_run(h3_text)
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_NORMAL
            i += 1
            continue

        h2_text = _match_h2(line)
        if h2_text:
            p = doc.add_paragraph(style="Heading 2")
            p.add_run(h2_text)
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_NORMAL
            i += 1
            continue

        h1_text = _match_h1(line)
        if h1_text:
            p = doc.add_paragraph(style="Heading 1")
            p.add_run(h1_text)
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_NORMAL
            i += 1
            continue

        # --- Плейсхолдер для рисунка ---
        ph = _handle_placeholder(line, image_map, images_dir)
        if ph is not None:
            fig_num, img_path = ph
            if img_path:
                insert_image(doc, img_path)
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
                img_file = image_map.get(fig_num) or image_map.get(int(fig_num))
                if img_file:
                    run = p.add_run(f"[Изображение не найдено: {img_file}]")
                    run.font.color.rgb = RGBColor(255, 0, 0)
                else:
                    run = p.add_run(line)
                    run.font.color.rgb = RGBColor(128, 128, 128)
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_NORMAL
            i += 1
            continue

        # --- Подпись рисунка ---
        m = RE_FIGURE_CAPTION.match(line)
        if m:
            p = doc.add_paragraph(style="FigureCaption")
            p.add_run(line)
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_NORMAL
            # Пустая строка после рисунка
            empty_p = doc.add_paragraph()
            empty_p.paragraph_format.space_before = Pt(0)
            empty_p.paragraph_format.space_after = Pt(0)
            empty_p.paragraph_format.line_spacing = LINE_SPACING
            empty_p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
            _set_contextual_spacing(empty_p)
            i += 1
            continue

        # --- Подпись таблицы ---
        m = RE_TABLE_CAPTION.match(line)
        if m:
            p = doc.add_paragraph(style="TableCaption")
            p.add_run(line)
            for run in p.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_NORMAL
            i += 1
            continue

        # --- Обычный абзац ---
        p = doc.add_paragraph(style="Normal")
        _set_contextual_spacing(p)
        add_formatted_run(p, line)
        i += 1


# ---------------------------------------------------------------------------
# Обработка отчёта в режиме шаблона-пустышки
# ---------------------------------------------------------------------------
def _insert_section_content(doc, heading_para, lines, heading_ref, image_map, images_dir):
    """Вставить содержимое секции после heading_para.

    Args:
        doc: Document
        heading_para: параграф-заголовок секции в шаблоне (вставляем после него)
        lines: строки содержимого секции
        heading_ref: эталонный заголовок для копирования стиля подзаголовков
        image_map: маппинг номер рисунка -> файл
        images_dir: директория с изображениями
    """
    last = heading_para
    in_code_block = False
    i = 0

    while i < len(lines):
        line = lines[i]

        # Блоки кода (```) — пропускаем целиком
        if RE_CODE_FENCE.match(line.strip()):
            in_code_block = not in_code_block
            i += 1
            continue
        if in_code_block:
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        # --- Markdown-таблица ---
        if RE_TABLE_ROW.match(line):
            table_lines, end_idx = _collect_table_lines(lines, i)
            headers, rows = _parse_markdown_table(table_lines)
            if headers and rows is not None:
                table = _create_word_table(doc, headers, rows, ref_para=last)
                # Параграф-якорь после таблицы
                dummy_p = OxmlElement('w:p')
                table._tbl.addnext(dummy_p)
                parent = doc._body if hasattr(doc, '_body') else doc.element.body
                last = Paragraph(dummy_p, parent)
            i = end_idx
            continue

        # --- Списки ---
        m = RE_LIST_BULLET.match(line)
        if m:
            last = _add_list_paragraph(doc, m.group(1), ref_para=last)
            i += 1
            continue

        m = RE_LIST_NUMBERED.match(line)
        if m:
            last = _add_list_paragraph(doc, m.group(2), number=m.group(1), ref_para=last)
            i += 1
            continue

        # --- Заголовки (H3, H2) ---

        h3_text = _match_h3(line)
        if h3_text:
            new_p = _insert_para_after(last, doc)
            if heading_ref:
                new_p.style = heading_ref.style
                _clone_paragraph_format(heading_ref, new_p)
            run = new_p.add_run(h3_text)
            if heading_ref and heading_ref.runs:
                _clone_run_format(heading_ref.runs[0], run)
            else:
                run.font.name = FONT_NAME
                run.font.bold = True
            run.font.size = FONT_SIZE_NORMAL
            last = new_p
            i += 1
            continue

        h2_text = _match_h2(line)
        if h2_text:
            new_p = _insert_para_after(last, doc)
            if heading_ref:
                new_p.style = heading_ref.style
                _clone_paragraph_format(heading_ref, new_p)
            run = new_p.add_run(h2_text)
            if heading_ref and heading_ref.runs:
                _clone_run_format(heading_ref.runs[0], run)
            else:
                run.font.name = FONT_NAME
                run.font.bold = True
            # Все шрифты строго 14 пт
            run.font.size = FONT_SIZE_NORMAL
            last = new_p
            i += 1
            continue

        # --- Плейсхолдер для рисунка ---
        ph = _handle_placeholder(line, image_map, images_dir)
        if ph is not None:
            fig_num, img_path = ph
            if img_path:
                last = _insert_image_at(doc, last, img_path)
            else:
                new_p = _insert_para_after(last, doc)
                new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                new_p.paragraph_format.left_indent = Cm(0)
                new_p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
                new_p.paragraph_format.line_spacing = LINE_SPACING
                _set_contextual_spacing(new_p)
                img_file = image_map.get(fig_num) or image_map.get(int(fig_num))
                if img_file:
                    run = new_p.add_run(f"[Изображение не найдено: {img_file}]")
                    run.font.color.rgb = RGBColor(255, 0, 0)
                else:
                    run = new_p.add_run(line)
                    run.font.color.rgb = RGBColor(128, 128, 128)
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_NORMAL
                last = new_p
            i += 1
            continue

        # --- Подпись рисунка ---
        m = RE_FIGURE_CAPTION.match(line)
        if m:
            new_p = _insert_para_after(last, doc)
            new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            new_p.paragraph_format.left_indent = Cm(0)
            new_p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
            new_p.paragraph_format.space_before = Pt(6)
            new_p.paragraph_format.space_after = Pt(0)
            new_p.paragraph_format.line_spacing = LINE_SPACING
            _set_contextual_spacing(new_p)
            run = new_p.add_run(line)
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE_NORMAL
            last = new_p
            # Пустая строка после рисунка
            empty_p = _insert_para_after(last, doc)
            empty_p.paragraph_format.left_indent = Cm(0)
            empty_p.paragraph_format.space_before = Pt(0)
            empty_p.paragraph_format.space_after = Pt(0)
            empty_p.paragraph_format.line_spacing = LINE_SPACING
            empty_p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
            _set_contextual_spacing(empty_p)
            last = empty_p
            i += 1
            continue

        # --- Подпись таблицы ---
        m = RE_TABLE_CAPTION.match(line)
        if m:
            new_p = _insert_para_after(last, doc)
            new_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            new_p.paragraph_format.left_indent = Cm(0)
            new_p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
            new_p.paragraph_format.space_before = Pt(0)
            new_p.paragraph_format.space_after = Pt(6)
            new_p.paragraph_format.line_spacing = LINE_SPACING
            _set_contextual_spacing(new_p)
            run = new_p.add_run(line)
            run.font.name = FONT_NAME
            run.font.size = FONT_SIZE_NORMAL
            last = new_p
            i += 1
            continue

        # --- Обычный абзац ---
        new_p = _insert_para_after(last, doc)
        new_p.paragraph_format.left_indent = Cm(0)
        new_p.paragraph_format.line_spacing = LINE_SPACING
        new_p.paragraph_format.first_line_indent = FIRST_LINE_INDENT
        new_p.paragraph_format.space_before = Pt(0)
        new_p.paragraph_format.space_after = Pt(0)
        new_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _set_contextual_spacing(new_p)
        add_formatted_run(new_p, line)
        last = new_p
        i += 1


def process_report_with_template(doc, report_text, image_map, images_dir):
    """Обработать report.md и вставить содержимое в шаблон-пустышку.

    Шаблон уже содержит титульный лист и заголовки секций.
    Стили и форматирование шаблона не изменяются.
    Подзаголовки создаются с тем же стилем, что и основные заголовки шаблона.
    """
    # Найти секции в шаблоне
    template_sections = find_template_sections(doc)

    if not template_sections:
        print(
            "Предупреждение: секции не найдены в шаблоне. "
            "Содержимое будет добавлено в конец документа.",
            file=sys.stderr,
        )
        setup_styles(doc)
        process_report(doc, report_text, image_map, images_dir)
        return

    print(f"Найдены секции в шаблоне: {', '.join(template_sections.keys())}")

    # Эталонный заголовок для копирования стиля подзаголовков
    heading_ref = None
    for key in SECTION_ORDER:
        if key in template_sections:
            heading_ref = template_sections[key]
            break

    # Разбить report.md на секции
    report_sections = parse_report_into_sections(report_text)

    # Вставить содержимое каждой секции
    inserted_count = 0
    for section_key, content_lines in report_sections:
        if section_key in template_sections:
            target_para = template_sections[section_key]
            non_empty = [l for l in content_lines if l.strip()]
            _insert_section_content(
                doc, target_para, content_lines,
                heading_ref, image_map, images_dir
            )
            inserted_count += len(non_empty)
            print(f"  Секция '{section_key}': вставлено {len(non_empty)} элементов")
        else:
            print(
                f"  Предупреждение: секция '{section_key}' не найдена в шаблоне — пропущена.",
                file=sys.stderr,
            )

    print(f"Всего вставлено элементов: {inserted_count}")


# ---------------------------------------------------------------------------
# Основная логика
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(
        description="Генератор ГОСТ-совместимого Word-документа из report.md"
    )
    parser.add_argument("--config", required=True, help="Путь к config.json")
    parser.add_argument("--report", required=True, help="Путь к report.md")
    parser.add_argument("--output", required=True, help="Путь к выходному файлу .docx")
    parser.add_argument("--images", default=None, help="Директория с изображениями")
    parser.add_argument("--image-map", default=None, help="Путь к image_map.json")
    parser.add_argument("--template", default=None,
                        help="Шаблон-пустышка .docx с титульным листом и заголовками")
    parser.add_argument("--no-title-page", action="store_true",
                        help="Не создавать титульный лист (режим без шаблона)")

    args = parser.parse_args()

    # Загрузить config
    config_path = Path(args.config)
    if not config_path.exists():
        print(f"Файл конфигурации не найден: {config_path}", file=sys.stderr)
        sys.exit(1)

    with open(config_path, encoding="utf-8") as f:
        config = json.load(f)

    # Загрузить report.md
    report_path = Path(args.report)
    if not report_path.exists():
        print(f"Файл отчёта не найден: {report_path}", file=sys.stderr)
        sys.exit(1)

    report_text = report_path.read_text(encoding="utf-8")

    # Загрузить image_map
    image_map = {}
    if args.image_map:
        map_path = Path(args.image_map)
        if map_path.exists():
            with open(map_path, encoding="utf-8") as f:
                image_map = json.load(f)

    images_dir = Path(args.images) if args.images else None

    # Создать / открыть документ
    if args.template:
        template_path = Path(args.template)
        if template_path.exists():
            doc = Document(str(template_path))
            # Режим шаблона: не трогаем стили, титульный лист, заголовки секций
            process_report_with_template(doc, report_text, image_map, images_dir)
        else:
            print(
                f"Шаблон не найден: {template_path}. Создаю документ с нуля.",
                file=sys.stderr,
            )
            doc = Document()
            setup_styles(doc)
            setup_page(doc)
            if not args.no_title_page:
                add_title_page(doc, config)
            process_report(doc, report_text, image_map, images_dir)
    else:
        doc = Document()
        setup_styles(doc)
        setup_page(doc)
        if not args.no_title_page:
            add_title_page(doc, config)
        process_report(doc, report_text, image_map, images_dir)

    # Сохранить
    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Документ сохранён: {output_path}")


if __name__ == "__main__":
    main()
