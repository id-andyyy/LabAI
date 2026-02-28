#!/usr/bin/env python3
"""Парсер документов лабораторных работ.

Извлекает текст из PDF, DOCX, TXT и MD файлов.

Режимы:
  - Обычный (по умолчанию): извлекает только текст → stdout.
  - Полный (--full): извлекает текст + изображения. Текст выводится в stdout
    в формате Markdown со ссылками на извлечённые изображения.
    Изображения сохраняются в директорию, указанную через --images-dir.
"""

import sys
import os
from pathlib import Path


def parse_pdf(file_path: Path) -> str:
    """Извлечь текст из PDF-файла."""
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n\n".join(pages)


def parse_pdf_full(file_path: Path, images_dir: Path) -> str:
    """Извлечь текст и изображения из PDF-файла.

    Возвращает Markdown с текстом и ссылками на извлечённые изображения.
    Изображения сохраняются в images_dir.
    """
    from pypdf import PdfReader

    reader = PdfReader(str(file_path))
    images_dir.mkdir(parents=True, exist_ok=True)

    result_parts = []
    image_counter = 0

    for page_num, page in enumerate(reader.pages, start=1):
        # Извлечь текст страницы
        text = page.extract_text()
        if text:
            result_parts.append(text)

        # Извлечь изображения страницы
        if hasattr(page, "images") and page.images:
            for image in page.images:
                image_counter += 1
                # Определить расширение по имени
                ext = Path(image.name).suffix if image.name else ".png"
                if not ext:
                    ext = ".png"
                image_filename = f"img_{image_counter:03d}{ext}"
                image_path = images_dir / image_filename

                with open(image_path, "wb") as f:
                    f.write(image.data)

                rel_path = os.path.relpath(image_path)
                result_parts.append(
                    f"![Изображение {image_counter} (стр. {page_num})]({rel_path})"
                )

    return "\n\n".join(result_parts)


def parse_docx(file_path: Path) -> str:
    """Извлечь текст из DOCX-файла."""
    from docx import Document

    doc = Document(str(file_path))
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))

    return "\n\n".join(parts)


def parse_docx_full(file_path: Path, images_dir: Path) -> str:
    """Извлечь текст и изображения из DOCX-файла.

    Возвращает Markdown с текстом и ссылками на извлечённые изображения.
    Изображения сохраняются в images_dir.
    """
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    doc = Document(str(file_path))
    images_dir.mkdir(parents=True, exist_ok=True)

    # Извлечь все изображения из документа
    image_counter = 0
    image_rels = {}  # rel_id -> relative_path

    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            image_counter += 1
            content_type = rel.target_part.content_type
            ext_map = {
                "image/png": ".png",
                "image/jpeg": ".jpg",
                "image/gif": ".gif",
                "image/bmp": ".bmp",
                "image/tiff": ".tiff",
                "image/svg+xml": ".svg",
                "image/x-emf": ".emf",
                "image/x-wmf": ".wmf",
            }
            ext = ext_map.get(content_type, ".png")
            image_filename = f"img_{image_counter:03d}{ext}"
            image_path = images_dir / image_filename

            with open(image_path, "wb") as f:
                f.write(rel.target_part.blob)

            rel_path = os.path.relpath(image_path)
            image_rels[rel_id] = rel_path

    # Собрать текст с информацией об изображениях
    parts = []
    image_insert_counter = 0

    for para in doc.paragraphs:
        text = para.text.strip()

        # Проверить, есть ли в параграфе изображения (inline shapes)
        has_images = False
        for run in para.runs:
            if run._element.findall(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing"
            ) or run._element.findall(
                ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict"
            ):
                has_images = True

        # Проверить через XML наличие blip (ссылок на изображения)
        blips = para._element.findall(
            ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
        )
        for blip in blips:
            embed = blip.get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
            if embed and embed in image_rels:
                image_insert_counter += 1
                parts.append(
                    f"![Изображение {image_insert_counter}]({image_rels[embed]})"
                )

        if text:
            parts.append(text)

    # Извлечь текст из таблиц
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))

    # Если были изображения, но не удалось привязать к параграфам —
    # добавить оставшиеся в конец
    referenced = set()
    for part in parts:
        for rel_id, rel_path in image_rels.items():
            if rel_path in part:
                referenced.add(rel_id)

    unreferenced_counter = image_insert_counter
    for rel_id, rel_path in image_rels.items():
        if rel_id not in referenced:
            unreferenced_counter += 1
            parts.append(
                f"![Изображение {unreferenced_counter} (не привязано к тексту)]({rel_path})"
            )

    return "\n\n".join(parts)


def parse_text(file_path: Path) -> str:
    """Прочитать текстовый файл (TXT/MD)."""
    return file_path.read_text(encoding="utf-8")


PARSERS = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".txt": parse_text,
    ".md": parse_text,
}

FULL_PARSERS = {
    ".pdf": parse_pdf_full,
    ".docx": parse_docx_full,
    ".txt": lambda fp, _: parse_text(fp),
    ".md": lambda fp, _: parse_text(fp),
}


def main():
    import argparse

    parser = argparse.ArgumentParser(
        description="Парсер документов лабораторных работ"
    )
    parser.add_argument("file", help="Путь к файлу задания")
    parser.add_argument(
        "--full",
        action="store_true",
        help="Полный парсинг: текст + изображения (Markdown с ссылками на картинки)",
    )
    parser.add_argument(
        "--images-dir",
        default=".claude/lab/source-images",
        help="Директория для сохранения извлечённых изображений (по умолчанию: .claude/lab/source-images)",
    )

    args = parser.parse_args()
    file_path = Path(args.file)

    if not file_path.exists():
        print(f"Файл не найден: {file_path}", file=sys.stderr)
        sys.exit(1)

    ext = file_path.suffix.lower()

    if args.full:
        full_parser = FULL_PARSERS.get(ext)
        if full_parser is None:
            supported = ", ".join(FULL_PARSERS.keys())
            print(
                f"Неподдерживаемый формат: {ext}. Поддерживаются: {supported}",
                file=sys.stderr,
            )
            sys.exit(1)

        try:
            images_dir = Path(args.images_dir)
            text = full_parser(file_path, images_dir)
            print(text)
        except Exception as e:
            print(f"Ошибка при обработке файла {file_path}: {e}", file=sys.stderr)
            sys.exit(1)
    else:
        text_parser = PARSERS.get(ext)
        if text_parser is None:
            supported = ", ".join(PARSERS.keys())
            print(
                f"Неподдерживаемый формат: {ext}. Поддерживаются: {supported}",
                file=sys.stderr,
            )
            sys.exit(1)

        try:
            text = text_parser(file_path)
            print(text)
        except Exception as e:
            print(f"Ошибка при обработке файла {file_path}: {e}", file=sys.stderr)
            sys.exit(1)


if __name__ == "__main__":
    main()
